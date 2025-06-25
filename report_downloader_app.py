# report_downloader_app.py

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import configparser
import os
import sys
import datetime
from typing import Dict, Any, Optional, List

# --- Dependency Imports ---
try:
    import pyodbc
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    # This logic has been simplified to be more robust and avoid syntax errors.
    try:
        missing_module_name = str(e).split("'")[1]
        error_message = f"A required library is missing: '{missing_module_name}'.\n\nPlease install it by running:\npip install {missing_module_name}"
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror("Missing Dependency", error_message)
    except (tk.TclError, IndexError):
        print(f"FATAL: Missing a required library. Original error: {e}", file=sys.stderr)
    sys.exit(1)

# --- Static Configuration ---
try:
    script_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
except NameError:
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))

CONFIG_FILE_NAME = "config.ini"
CONFIG_FILE_PATH = os.path.join(script_dir, CONFIG_FILE_NAME)


# --- DOCX Generation Logic (Corrected Formatting) ---

def _add_main_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)

def _add_section_heading(doc: Document, text: str):
    doc.add_heading(text, level=2)

def add_email_body_section(doc: Document, data: Dict[str, Any]):
    _add_section_heading(doc, "Performance Summary")
    header = data.get("report_header", {})
    p = doc.add_paragraph()
    p.add_run("Agent: ").bold = True
    p.add_run(f"{header.get('AgentName', 'N/A')}\n")
    p.add_run("Period: ").bold = True
    p.add_run(f"{header.get('AnalysisPeriodNote', 'N/A')}\n")
    p.add_run("Calls Analyzed: ").bold = True
    p.add_run(f"{header.get('NumberOfReportsSuccessfullyAnalyzed', 'N/A')}")
    doc.add_paragraph()

def add_combined_analysis_section(doc: Document, data: Dict[str, Any]):
    _add_section_heading(doc, "Combined Analysis Report")
    summary = data.get("qualitative_summary_and_coaching_plan", {})
    if items := summary.get("overall_strengths_observed"):
        doc.add_heading("Key Strengths", level=3)
        for i in items:
            doc.add_paragraph(i, style='ListBullet')
    if items := summary.get("overall_areas_for_development"):
        doc.add_heading("Areas for Development", level=3)
        for i in items:
            doc.add_paragraph(i, style='ListBullet')
    if items := summary.get("consolidated_coaching_focus"):
        doc.add_heading("Coaching Plan", level=3)
        for item in items:
            doc.add_heading(item.get('area', 'Focus Area'), level=4)
            if actions := item.get("specific_actions"):
                for a in actions:
                    doc.add_paragraph(a, style='ListNumber')
    if details := data.get("detailed_quality_point_analysis"):
        doc.add_heading("Detailed Quality Point Analysis", level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Quality Point"
        hdr[1].text = "Trend Observation"
        for item in details:
            row = table.add_row().cells
            row[0].text = item.get("QualityPointText", "N/A")
            row[1].text = item.get("TrendObservation", "N/A")

def add_individual_qas_section(doc: Document, data: List[Dict[str, Any]]):
    _add_section_heading(doc, "Individual Call Analyses")
    for i, call_data in enumerate(data):
        summary = call_data.get('summary', {})
        doc.add_heading(f"Analysis for: {summary.get('OriginalAudioFileName', f'Call {i+1}')}", level=3)
        p = doc.add_paragraph()
        p.add_run("Client: ").bold = True
        p.add_run(f"{summary.get('ClientName', 'N/A')}\n")
        p.add_run("Ticket #: ").bold = True
        p.add_run(f"{summary.get('TicketNumber', 'N/A')}\n")

        if items := call_data.get('evaluation_items', []):
            doc.add_heading("Evaluation Findings", level=4)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Quality Point"
            hdr[1].text = "Finding"
            hdr[2].text = "Explanation"
            for item in items:
                row = table.add_row().cells
                row[0].text = item.get("QualityPointText", "N/A")
                row[1].text = item.get("Finding", "N/A")
                row[2].text = item.get("ExplanationSnippets", "")
        if i < len(data) - 1:
            doc.add_page_break()


class ReportDownloaderApp:
    def __init__(self, master: tk.Tk):
        self.master = master
        master.title("QA Report Downloader (Enhanced)")
        master.geometry("650x550")

        db_settings = self._load_db_config()
        if not db_settings:
             master.destroy()
             return

        self.conn = self._get_db_connection(db_settings)
        if not self.conn:
            master.destroy()
            return

        self.agents: Dict[str, int] = {}
        self.analysis_details: Dict[str, Dict] = {}
        self._setup_ui()
        self.populate_agent_list()

    def _load_db_config(self) -> Optional[Dict[str, str]]:
        if not os.path.exists(CONFIG_FILE_PATH):
            messagebox.showerror("Configuration Error", f"Config file '{CONFIG_FILE_NAME}' not found.")
            return None
        try:
            config = configparser.ConfigParser()
            config.read(CONFIG_FILE_PATH)
            return dict(config.items('Database'))
        except (configparser.Error, KeyError) as e:
            messagebox.showerror("Configuration Error", f"Missing or invalid [Database] section in config.ini: {e}")
            return None

    def _get_db_connection(self, db_cfg: Dict[str, str]) -> Optional[pyodbc.Connection]:
        try:
            conn_str = f"DRIVER={{ODBC Driver 17 for SQL Server}};SERVER={db_cfg['server']};DATABASE={db_cfg['database']};UID={db_cfg['user']};PWD={db_cfg['password']};"
            return pyodbc.connect(conn_str)
        except pyodbc.Error as e:
            messagebox.showerror("Database Connection Error", f"Could not connect to the database:\n{e}")
            return None

    def _setup_ui(self):
        ttk.Label(self.master, text="1. Select an Agent:", font=("Segoe UI", 10, "bold")).pack(pady=(10,2), padx=10, anchor="w")
        self.agent_listbox = tk.Listbox(self.master, exportselection=False, height=8)
        self.agent_listbox.pack(fill="x", expand=True, padx=10)
        self.agent_listbox.bind("<<ListboxSelect>>", self.on_agent_select)

        ttk.Label(self.master, text="2. Select a Combined Report:", font=("Segoe UI", 10, "bold")).pack(pady=(10,2), padx=10, anchor="w")
        self.date_listbox = tk.Listbox(self.master, exportselection=False, height=8)
        self.date_listbox.pack(fill="x", expand=True, padx=10)

        ttk.Label(self.master, text="3. Choose Content to Include:", font=("Segoe UI", 10, "bold")).pack(pady=(10,5), padx=10, anchor="w")
        self.include_combined = tk.BooleanVar(value=True)
        self.include_individual = tk.BooleanVar(value=True)
        ttk.Checkbutton(self.master, text="Combined Analysis Report", variable=self.include_combined).pack(anchor="w", padx=20)
        ttk.Checkbutton(self.master, text="Associated Individual Call Details", variable=self.include_individual).pack(anchor="w", padx=20)

        ttk.Button(self.master, text="Download Selected Report as .DOCX", command=self.on_download_click).pack(pady=20, padx=10, fill="x", ipady=5)

    def populate_agent_list(self):
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT AgentID, AgentName FROM Agents ORDER BY AgentName;")
            for row in cursor.fetchall():
                self.agents[row.AgentName] = row.AgentID
                self.agent_listbox.insert(tk.END, row.AgentName)
        except pyodbc.Error as e:
            messagebox.showerror("Error", f"Failed to fetch agent list:\n{e}")

    def on_agent_select(self, event=None):
        if not (selections := self.agent_listbox.curselection()):
            return
        self.date_listbox.delete(0, tk.END)
        self.analysis_details.clear()
        agent_id = self.agents.get(self.agent_listbox.get(selections[0]))
        if agent_id:
            try:
                cursor = self.conn.cursor()
                sql = "SELECT CombinedAnalysisID, ProcessingDateTime FROM CombinedAnalyses WHERE AgentID = ? ORDER BY ProcessingDateTime DESC;"
                for row in cursor.execute(sql, agent_id).fetchall():
                    date_str = row.ProcessingDateTime.strftime("%Y-%m-%d %I:%M %p") if row.ProcessingDateTime else f"Report ID {row.CombinedAnalysisID}"
                    self.analysis_details[date_str] = {"id": row.CombinedAnalysisID, "timestamp": row.ProcessingDateTime}
                    self.date_listbox.insert(tk.END, date_str)
            except pyodbc.Error as e:
                messagebox.showerror("Error", f"Failed to fetch reports:\n{e}")

    def on_download_click(self):
        if not self.agent_listbox.curselection() or not self.date_listbox.curselection():
            messagebox.showwarning("Selection Required", "Please select an agent and a report.")
            return
        if not self.include_combined.get() and not self.include_individual.get():
            messagebox.showwarning("Selection Required", "Please check at least one content type to include.")
            return

        selected_date_str = self.date_listbox.get(self.date_listbox.curselection())
        details = self.analysis_details.get(selected_date_str)

        if self.include_individual.get() and (not details or not details.get('timestamp')):
            messagebox.showerror("Data Compatibility Error", "Cannot download individual details for this report as it lacks the required 'ProcessingDateTime' link.\nPlease re-import the data for this period.")
            return

        agent_name = self.agent_listbox.get(self.agent_listbox.curselection())
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")], title="Save Report As...", initialfile=f"QA_Report_Full_{agent_name}_{selected_date_str.split(' ')[0]}.docx")
        if not save_path: return

        try:
            doc = Document()
            _add_main_title(doc, f"Quality Assurance Report for {agent_name}")
            combined_data = self.fetch_combined_analysis_data(details['id']) if self.include_combined.get() else None
            individual_data = self.fetch_individual_qas_data(self.agents[agent_name], details['timestamp']) if self.include_individual.get() and details.get('timestamp') else None

            if not combined_data and not individual_data:
                messagebox.showerror("Data Error", "No data was selected or found for report generation.")
                return

            if combined_data:
                add_email_body_section(doc, combined_data)
                add_combined_analysis_section(doc, combined_data)
            if individual_data:
                if combined_data: doc.add_page_break()
                add_individual_qas_section(doc, individual_data)
                
            doc.save(save_path)

            # --- MODIFICATION START ---
            # Show a success message and then attempt to open the file.
            messagebox.showinfo("Success", f"Report successfully saved to:\n{save_path}")
            try:
                os.startfile(save_path)
            except AttributeError:
                # os.startfile() is only available on Windows.
                # You can add alternatives for macOS and Linux here if needed.
                # For example, using the 'subprocess' module:
                # import subprocess
                # subprocess.call(['open', save_path])  # macOS
                # subprocess.call(['xdg-open', save_path]) # Linux
                messagebox.showwarning("Cannot Open File", "File saved, but could not be opened automatically on this OS.")
            except Exception as e:
                # Handle other potential errors, like file not found (unlikely) or no associated application.
                messagebox.showwarning("Could Not Open File", f"The report was saved, but an error occurred while trying to open it:\n{e}")
            # --- MODIFICATION END ---

        except Exception as e:
            messagebox.showerror("Report Generation Failed", f"An unexpected error occurred:\n{e}")

    def fetch_combined_analysis_data(self, analysis_id: int) -> Optional[Dict]:
        try:
            cursor = self.conn.cursor()
            main_row = cursor.execute("SELECT * FROM CombinedAnalyses c JOIN Agents a ON c.AgentID = a.AgentID WHERE c.CombinedAnalysisID = ?;", analysis_id).fetchone()
            if not main_row: return None

            report = {
                "report_header": dict(zip([c[0] for c in main_row.cursor_description], main_row)),
                "qualitative_summary_and_coaching_plan": {},
                "detailed_quality_point_analysis": []
            }
            qscp = report["qualitative_summary_and_coaching_plan"]
            qscp["overall_strengths_observed"] = [r.StrengthText for r in cursor.execute("SELECT StrengthText FROM CombinedAnalysisStrengths WHERE CombinedAnalysisID = ?", analysis_id)]
            qscp["overall_areas_for_development"] = [r.DevelopmentAreaText for r in cursor.execute("SELECT DevelopmentAreaText FROM CombinedAnalysisDevelopmentAreas WHERE CombinedAnalysisID = ?", analysis_id)]

            focus_items = []
            for focus_row in cursor.execute("SELECT CoachingFocusID, AreaText FROM CombinedAnalysisCoachingFocus WHERE CombinedAnalysisID = ?", analysis_id):
                actions = [r.ActionText for r in cursor.execute("SELECT ActionText FROM CombinedAnalysisCoachingActions WHERE CoachingFocusID = ?", focus_row.CoachingFocusID)]
                focus_items.append({"area": focus_row.AreaText, "specific_actions": actions})
            qscp["consolidated_coaching_focus"] = focus_items

            report["detailed_quality_point_analysis"] = [dict(zip([c[0] for c in r.cursor_description], r)) for r in cursor.execute("SELECT qp.QualityPointText, d.TrendObservation FROM CombinedAnalysisQualityPointDetails d JOIN QualityPointsMaster qp ON d.QualityPointID = qp.QualityPointID WHERE d.CombinedAnalysisID = ?;", analysis_id)]
            return report
        except pyodbc.Error as e:
            messagebox.showerror("Database Error", f"Could not fetch combined report data:\n{e}")
            return None

    def fetch_individual_qas_data(self, agent_id: int, ts: datetime.datetime) -> List[Dict]:
        try:
            cursor = self.conn.cursor()
            results = []
            main_rows = cursor.execute("SELECT * FROM IndividualCallAnalyses WHERE AgentID = ? AND ProcessingDateTime = ?;", agent_id, ts).fetchall()
            for row in main_rows:
                items_sql = "SELECT qpm.QualityPointText, iei.Finding, iei.ExplanationSnippets FROM IndividualEvaluationItems iei JOIN QualityPointsMaster qpm ON iei.QualityPointID = qpm.QualityPointID WHERE iei.AnalysisID = ?;"
                item_rows = cursor.execute(items_sql, row.AnalysisID).fetchall()
                results.append({
                    "summary": dict(zip([c[0] for c in row.cursor_description], row)),
                    "evaluation_items": [dict(zip([c[0] for c in i.cursor_description], i)) for i in item_rows]
                })
            return results
        except pyodbc.Error as e:
            messagebox.showerror("Database Error", f"Could not fetch individual QA data:\n{e}")
            return []

    def on_closing(self):
        if self.conn:
            self.conn.close()
        self.master.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = ReportDownloaderApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()