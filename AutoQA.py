# AutoQA.py

# --- Standard Library Imports ---
import os
import sys
import datetime
import glob
import re
import logging
import argparse
import configparser
import traceback
import smtplib
import json
import time
import shutil
import base64
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from typing import Dict, Any, Optional, List, Tuple

# --- Third-Party Library Imports ---
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    print("FATAL ERROR: The 'python-docx' library is not installed. Please install it with 'pip install python-docx'", file=sys.stderr)
    sys.exit(1)

try:
    import google.generativeai as genai
    from google.api_core import exceptions as google_api_exceptions
except ImportError:
    print("FATAL ERROR: The 'google-generativeai' library is not installed. Please install it with 'pip install google-generativeai'", file=sys.stderr)
    sys.exit(1)

# --- Determine Script Directory ---
try:
    script_dir = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
except NameError:
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))

# --- Static Configuration ---
CONFIG_FILE_NAME = "config.ini"
CONFIG_FILE_PATH = os.path.join(script_dir, CONFIG_FILE_NAME)
EXT_LIST_FILE_PATH = os.path.join(script_dir, "ExtList.data")

# --- Global Logger Instance ---
logger = logging.getLogger("AutoQA")


# --- Utility and Helper Functions ---

def setup_logger(log_dir: str, is_debug: bool):
    """Configures the global logger."""
    os.makedirs(log_dir, exist_ok=True)
    log_file = os.path.join(log_dir, f"AutoQA_{datetime.datetime.now():%Y%m%d_%H%M%S}.log")
    
    logger.setLevel(logging.DEBUG)
    if logger.hasHandlers():
        logger.handlers.clear()

    fh = logging.FileHandler(log_file, encoding='utf-8')
    fh.setLevel(logging.DEBUG)
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] - %(message)s')
    fh.setFormatter(formatter)
    logger.addHandler(fh)

    ch = logging.StreamHandler(sys.stdout)
    ch.setLevel(logging.DEBUG if is_debug else logging.INFO)
    ch_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    ch.setFormatter(ch_formatter)
    logger.addHandler(ch)
    
    logger.info(f"Logger initialized. Debug mode: {is_debug}. Log file: {log_file}")

def load_text_from_file(file_path: str, purpose: str) -> str:
    """Loads text content from a file, exiting on failure."""
    logger.info(f"Loading {purpose} from: {file_path}")
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        if not content.strip():
            logger.critical(f"{purpose.capitalize()} file '{file_path}' is empty.")
            sys.exit(1)
        return content
    except FileNotFoundError:
        logger.critical(f"Required file for {purpose} not found: '{file_path}'.")
        sys.exit(1)
    except Exception as e:
        logger.critical(f"Error reading {purpose} file '{file_path}': {e}", exc_info=True)
        sys.exit(1)

def save_text_to_file(file_path: str, content: str):
    """Saves text content to a file, creating directories if needed."""
    try:
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.debug(f"Successfully saved text to: {file_path}")
    except IOError as e:
        logger.error(f"File save failed for '{file_path}': {e}", exc_info=True)

def sanitize_filename(name: str) -> str:
    """Removes invalid characters from a string to make it a valid filename component."""
    return re.sub(r'[<>:"/\\|?*\x00-\x1F]', '_', name).strip()

def clean_ai_response(text: str) -> str:
    """Strips markdown code fences and extracts content between the first { and last }."""
    if text.lower().startswith("```json"):
        text = text[len("```json"):].strip()
    if text.lower().endswith("```"):
        text = text[:-len("```")].strip()
    
    try:
        first_brace = text.find('{')
        last_brace = text.rfind('}')
        if first_brace != -1 and last_brace > first_brace:
            return text[first_brace : last_brace + 1]
    except Exception:
        pass
    return text

def fetch_member_list() -> Dict[str, Dict[str, str]]:
    """
    Fetches the member list from ExtList.data.
    Handles an optional 4th column for a custom prompt file.
    """
    members = {}
    if not os.path.exists(EXT_LIST_FILE_PATH):
        logger.critical(f"Agent data file not found: '{EXT_LIST_FILE_PATH}'.")
        sys.exit(1)
    try:
        with open(EXT_LIST_FILE_PATH, 'r', encoding='utf-8') as f:
            for line in f:
                s_line = line.strip()
                if not s_line or s_line.startswith('#'): continue
                parts = s_line.split('\t')
                if len(parts) >= 3:
                    ext, name, email = parts[0].strip(), parts[1].strip(), parts[2].strip()
                    # NEW: Check for a 4th column specifying a custom prompt file
                    prompt_file = parts[3].strip() if len(parts) >= 4 and parts[3].strip() else None
                    if ext: members[ext] = {"full_name": name, "email": email, "prompt_file": prompt_file}
        logger.info(f"Fetched {len(members)} members from {os.path.basename(EXT_LIST_FILE_PATH)}.")
        return members
    except Exception as e:
        logger.critical(f"Error processing agent data file '{EXT_LIST_FILE_PATH}': {e}", exc_info=True)
        sys.exit(1)

def _convert_individual_json_to_html_string(json_data: Dict[str, Any], filename: str) -> str:
    """Converts a single QA JSON into a simple HTML string for AI consumption."""
    parts = [f"<h1>QA Analysis for: {filename.replace('_analysis.json','')}</h1>"]
    if summary := json_data.get("call_summary"):
        parts.append("<h2>I. Call Summary</h2><ul>")
        for key, value in summary.items():
            parts.append(f"<li><strong>{key.replace('_', ' ').title()}:</strong> {value}</li>")
        parts.append("</ul>")
    if evaluation := json_data.get("detailed_evaluation"):
        parts.append("<h2>II. Detailed Evaluation</h2><table border='1' style='border-collapse: collapse; width: 100%;'>")
        parts.append("<thead><tr><th>Quality Point</th><th>Finding</th><th>Explanation</th></tr></thead><tbody>")
        for item in evaluation:
            parts.append(f"<tr><td>{item.get('quality_point')}</td><td>{item.get('finding')}</td><td>{item.get('explanation_snippets')}</td></tr>")
        parts.append("</tbody></table>")
    if remarks := json_data.get("concluding_remarks"):
        parts.append("<h2>III. Concluding Remarks</h2>")
        if text := remarks.get("summary_positive_findings"): parts.append(f"<h3>Positive Findings</h3><p>{text}</p>")
        if text := remarks.get("summary_negative_findings"): parts.append(f"<h3>Negative Findings</h3><p>{text}</p>")
        if text := remarks.get("coaching_plan_for_growth"): parts.append(f"<h3>Coaching Plan</h3><p>{text}</p>")
    return f"<!DOCTYPE html><html><body>{''.join(parts)}</body></html>"

def _convert_combined_json_to_html_string(json_data: Dict[str, Any]) -> str:
    """Converts a combined analysis JSON into an HTML string."""
    parts = []
    if header := json_data.get("report_header"):
        parts.append(f"<h1>Performance Summary for {header.get('agent_name', 'N/A')}</h1>")
        parts.append(f"<p>Period: {header.get('analysis_period_note', 'N/A')}</p>")
    if summary := json_data.get("qualitative_summary_and_coaching_plan"):
        if items := summary.get("overall_strengths_observed"):
            parts.append("<h2>Key Strengths</h2><ul>"); [parts.append(f"<li>{item}</li>") for item in items]; parts.append("</ul>")
        if items := summary.get("overall_areas_for_development"):
            parts.append("<h2>Areas for Development</h2><ul>"); [parts.append(f"<li>{item}</li>") for item in items]; parts.append("</ul>")
        if items := summary.get("consolidated_coaching_focus"):
            parts.append("<h2>Coaching Plan</h2>")
            for item in items:
                parts.append(f"<h3>{item.get('area')}</h3>")
                if actions := item.get("specific_actions"):
                    parts.append("<ul>"); [parts.append(f"<li>{action}</li>") for action in actions]; parts.append("</ul>")
    return f"<!DOCTYPE html><html><body>{''.join(parts)}</body></html>"

def get_ai_model(api_key: str, model_name: str) -> genai.GenerativeModel:
    """Initializes and returns the Generative AI model."""
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        logger.info(f"AI Model '{model_name}' initialized successfully.")
        return model
    except Exception as e:
        logger.critical(f"Failed to initialize AI model '{model_name}': {e}", exc_info=True)
        sys.exit(1)

def call_ai_with_retry(model: genai.GenerativeModel, content: List[Any], log_prefix: str) -> Optional[genai.types.GenerateContentResponse]:
    """Calls the Generative AI model with a retry mechanism for specific errors."""
    for attempt in range(2):
        try:
            logger.info(f"{log_prefix}: AI call attempt {attempt + 1}...")
            response = model.generate_content(content)
            logger.info(f"{log_prefix}: AI call successful.")
            return response
        except (google_api_exceptions.InternalServerError, google_api_exceptions.ServiceUnavailable, google_api_exceptions.ResourceExhausted) as e:
            logger.warning(f"{log_prefix}: AI call failed with retriable error: {type(e).__name__}. Retrying in 5 seconds...")
            time.sleep(5)
        except Exception as e:
            logger.error(f"{log_prefix}: AI call failed with non-retriable error: {e}", exc_info=True)
            return None
    logger.error(f"{log_prefix}: AI call failed after all retry attempts.")
    return None

def send_email(smtp_cfg: Dict[str, str], email_cfg: Dict[str, Any], to_addr: str, subject: str, body: str, attachment_path: Optional[str] = None):
    """Constructs and sends an email using configured SMTP settings."""
    if not to_addr or '@' not in to_addr:
        logger.error(f"Invalid 'To' address for email with subject '{subject}'. Skipping send.")
        return
    try:
        from_addr = email_cfg['fromaddress']
        recipients = [to_addr] + email_cfg.get('ccaddresses', [])
        
        msg = MIMEMultipart('related')
        msg['Subject'] = subject
        msg['From'] = from_addr
        msg['To'] = to_addr
        if cc_list := email_cfg.get('ccaddresses'): msg['Cc'] = ', '.join(cc_list)
        if reply_to := email_cfg.get('replytoaddress'): msg['Reply-To'] = reply_to
        
        msg.attach(MIMEText(body, 'html', 'utf-8'))

        if attachment_path and os.path.exists(attachment_path):
            with open(attachment_path, "rb") as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(attachment_path))
            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
            msg.attach(part)
            logger.info(f"Attached '{os.path.basename(attachment_path)}' to email.")
        
        logger.info(f"Sending email to {recipients} via {smtp_cfg['server']}...")
        pwd = base64.b64decode(smtp_cfg['password_b64']).decode('utf-8')
        
        use_ssl = smtp_cfg.get('ssl', 'false').lower() == 'true'
        port = int(smtp_cfg.get('port', 587))
        server = smtplib.SMTP_SSL(smtp_cfg['server'], port) if use_ssl else smtplib.SMTP(smtp_cfg['server'], port)
        if not use_ssl and smtp_cfg.get('usestarttls', 'true').lower() == 'true':
            server.starttls()
        
        server.login(smtp_cfg['uid'], pwd)
        server.sendmail(from_addr, recipients, msg.as_string())
        server.quit()
        logger.info("Email sent successfully.")
    except Exception as e:
        logger.error(f"Failed to send email to '{to_addr}': {e}", exc_info=True)

def _apply_finding_color(run: Any, finding: str):
    if finding == "Positive": run.font.color.rgb = RGBColor(0x00, 0x64, 0x00)
    elif finding == "Negative": run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif finding == "Neutral": run.font.color.rgb = RGBColor(0x00, 0x00, 0x8B)
    run.bold = True

def generate_individual_docx(json_data: Dict, docx_path: str, agent_name: str, audio_filename: str, week_str: str):
    try:
        doc = Document(); doc.add_heading(f"Call Quality Assurance Report: {agent_name}", level=1)
        doc.add_paragraph(f"Audio File: {audio_filename}\nWeek Starting: {week_str}")
        if summary := json_data.get("call_summary"):
            doc.add_heading("I. Call Summary", level=2)
            for key, value in summary.items():
                p = doc.add_paragraph(); p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True; p.add_run(str(value))
        if evaluation := json_data.get("detailed_evaluation"):
            doc.add_heading("II. Detailed Evaluation", level=2)
            table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
            for i, name in enumerate(["Quality Point", "Finding", "Explanation & Snippets"]):
                table.cell(0, i).text = name; table.cell(0, i).paragraphs[0].runs[0].bold = True
            for item in evaluation:
                cells = table.add_row().cells; cells[0].text = item.get("quality_point", "")
                _apply_finding_color(cells[1].paragraphs[0].add_run(item.get("finding", "")), item.get("finding", ""))
                cells[2].text = item.get("explanation_snippets", "")
        if remarks := json_data.get("concluding_remarks"):
            doc.add_heading("III. Concluding Remarks", level=2)
            for key, heading in [("summary_positive_findings", "Positive"), ("summary_negative_findings", "Negative"), ("coaching_plan_for_growth", "Coaching")]:
                if text := remarks.get(key): doc.add_heading(heading, level=3); doc.add_paragraph(text)
        doc.save(docx_path)
        logger.info(f"Successfully generated DOCX: {os.path.basename(docx_path)}")
    except Exception as e: logger.error(f"Failed to generate individual DOCX '{docx_path}': {e}", exc_info=True)

def generate_combined_docx(json_data: Dict, docx_path: str, agent_name: str):
    try:
        doc = Document()
        header = json_data.get("report_header", {}); doc.add_heading(f"Performance Trend Analysis & Coaching Report: {agent_name}", level=1)
        doc.add_paragraph(f"Analysis Period: {header.get('analysis_period_note', 'N/A')}")
        if summary := json_data.get("qualitative_summary_and_coaching_plan"):
            if items := summary.get("overall_strengths_observed"):
                doc.add_heading("Key Strengths", level=2); [doc.add_paragraph(i, style='ListBullet') for i in items]
            if items := summary.get("overall_areas_for_development"):
                doc.add_heading("Areas for Development", level=2); [doc.add_paragraph(i, style='ListBullet') for i in items]
            if items := summary.get("consolidated_coaching_focus"):
                doc.add_heading("Coaching Plan", level=2)
                for item in items:
                    doc.add_heading(item.get('area', 'Focus Area'), level=3)
                    if actions := item.get("specific_actions"): [doc.add_paragraph(a, style='ListNumber') for a in actions]
        if details := json_data.get("detailed_quality_point_analysis"):
            doc.add_heading("Detailed Quality Point Analysis", level=2)
            table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
            hdr = table.rows[0].cells; hdr[0].text = "Quality Point"; hdr[1].text = "Trend Observation"
            for item in details:
                row = table.add_row().cells; row[0].text = item.get("quality_point", ""); row[1].text = item.get("trend_observation", "")
        doc.save(docx_path)
        logger.info(f"Successfully generated combined DOCX: {os.path.basename(docx_path)}")
    except Exception as e: logger.error(f"Failed to generate combined DOCX '{docx_path}': {e}", exc_info=True)

def process_single_audio(wav_path: str, output_dirs: Dict, model: genai.GenerativeModel, prompt: str, agent_name: str, week_start_str: str) -> Optional[Dict]:
    log_prefix = f"File '{os.path.basename(wav_path)}':"; sanitized_name = sanitize_filename(os.path.splitext(os.path.basename(wav_path))[0])
    json_path = os.path.join(output_dirs['json'], f"{sanitized_name}_analysis.json")
    docx_path = os.path.join(output_dirs['docx'], f"{sanitized_name}_analysis.docx")
    logger.info(f"Analyzing: {os.path.basename(wav_path)}"); uploaded_file = None
    try:
        uploaded_file = genai.upload_file(path=wav_path)
        response = call_ai_with_retry(model, [prompt, uploaded_file], log_prefix)
        if not response or not response.text: raise ValueError("AI returned no response.")
        cleaned_text = clean_ai_response(response.text); json_data = json.loads(cleaned_text)
        save_text_to_file(json_path, json.dumps(json_data, indent=2))
        generate_individual_docx(json_data, docx_path, agent_name, os.path.basename(wav_path), week_start_str)
        return json_data
    except Exception as e:
        logger.error(f"{log_prefix} Failed single audio processing: {e}", exc_info=True); return None
    finally:
        if uploaded_file:
            try: genai.delete_file(uploaded_file.name)
            except Exception as e: logger.warning(f"Could not delete AI service file {uploaded_file.name}: {e}")

def perform_combined_analysis(jsons: List[Dict], agent_name: str, reports_dir: str, model: genai.GenerativeModel, prompt: str) -> Optional[Dict]:
    logger.info(f"Starting combined analysis for {agent_name} using HTML conversion method...")
    temp_dir = os.path.join(reports_dir, "temp_html_combined"); uploaded_files = []
    try:
        os.makedirs(temp_dir, exist_ok=True)
        for i, data in enumerate(jsons):
            html_content = _convert_individual_json_to_html_string(data, f"report_{i+1}")
            html_path = os.path.join(temp_dir, f"report_{i+1}.html")
            save_text_to_file(html_path, html_content)
            uploaded_files.append(genai.upload_file(path=html_path))
        if not uploaded_files: logger.error("No HTML files were uploaded for combined analysis."); return None
        content = [prompt.replace("<AgentNamePlaceholder>", agent_name)] + uploaded_files
        response = call_ai_with_retry(model, content, f"Combined for {agent_name}")
        if not response or not response.text: raise ValueError("AI returned no response for combined analysis.")
        return json.loads(clean_ai_response(response.text))
    except Exception as e:
        logger.error(f"Error during combined analysis for {agent_name}: {e}", exc_info=True); return None
    finally:
        for uf in uploaded_files:
            try: genai.delete_file(uf.name)
            except Exception as e: logger.warning(f"Failed to delete temp HTML file {uf.name}: {e}")
        if os.path.exists(temp_dir):
            try: shutil.rmtree(temp_dir)
            except Exception as e: logger.warning(f"Failed to remove temp dir '{temp_dir}': {e}")

def generate_email_body(data: Dict, agent_name: str, model: genai.GenerativeModel, prompt: str) -> Optional[Dict]:
    logger.info(f"Generating email content for {agent_name}...")
    html_content = _convert_combined_json_to_html_string(data); temp_path = os.path.join(script_dir, "temp_email_input.html"); uploaded_file = None
    try:
        save_text_to_file(temp_path, html_content)
        uploaded_file = genai.upload_file(path=temp_path)
        content = [prompt.replace("<AgentNamePlaceholder>", agent_name), uploaded_file]
        response = call_ai_with_retry(model, content, f"EmailGen for {agent_name}")
        if not response or not response.text: raise ValueError("AI returned no response for email generation.")
        return json.loads(clean_ai_response(response.text))
    except Exception as e:
        logger.error(f"Error during email generation for {agent_name}: {e}", exc_info=True); return None
    finally:
        if uploaded_file:
            try: genai.delete_file(uploaded_file.name)
            except Exception as e: logger.warning(f"Failed to delete email temp file {uploaded_file.name}: {e}")
        if os.path.exists(temp_path):
            try: os.remove(temp_path)
            except Exception as e: logger.warning(f"Failed to remove temp file '{temp_path}': {e}")

def parse_arguments() -> Tuple[Optional[datetime.date], bool]:
    parser = argparse.ArgumentParser(description="Automated Phone QA analysis and reporting.")
    parser.add_argument("-D", "--date", type=str, help="Specify Sunday date (YYYY-MM-DD) to define processing week.")
    parser.add_argument('--debug', action='store_true', help='Enable detailed debug console messages.')
    args = parser.parse_args()
    date_override = None
    if args.date:
        try:
            date_override = datetime.datetime.strptime(args.date, '%Y-%m-%d').date()
            if date_override.weekday() != 6: print(f"WARNING: Date '{args.date}' is not a Sunday.", file=sys.stderr)
        except ValueError: print(f"ERROR: Invalid date format for --date: '{args.date}'. Ignoring.", file=sys.stderr)
    return date_override, args.debug

def main(config: configparser.ConfigParser, cli_date_override: Optional[datetime.date], is_debug_mode: bool):
    log_output_root = config.get('Paths', 'AutoQALogOutputRoot')
    setup_logger(log_output_root, is_debug_mode)
    
    try:
        source_root = config.get('Paths', 'ImporterSourceRoot')
        api_key = base64.b64decode(config.get('API', 'API_Key_B64')).decode('utf-8')
        model_name = config.get('API', 'ModelName')
        smtp_settings = dict(config.items('SMTP'))
        email_settings = dict(config.items('AutoQA Emails'))
        email_settings['ccaddresses'] = [addr.strip() for addr in email_settings.get('ccaddresses', '').split(';') if addr.strip()]
        
        prompt_dir = os.path.join(script_dir, "prompts")
        # MODIFIED: Load default prompts, custom ones will be loaded on-demand
        prompts = {
            'individual': load_text_from_file(os.path.join(prompt_dir, config.get('Prompts', 'IndividualPromptFile')), "Default Individual"),
            'combined': load_text_from_file(os.path.join(prompt_dir, config.get('Prompts', 'CombinedPromptFile')), "Combined"),
            'email': load_text_from_file(os.path.join(prompt_dir, config.get('Prompts', 'EmailPromptFile')), "Email")
        }
    except (configparser.Error, KeyError, base64.B64DecodeError) as e:
        logger.critical(f"FATAL: Failed to load settings from config.ini: {e}", exc_info=True); sys.exit(1)

    ai_model = get_ai_model(api_key, model_name)
    member_list = fetch_member_list()

    effective_date = cli_date_override or datetime.date.today()
    days_to_sunday = (effective_date.weekday() + 1) % 7
    week_start = effective_date - datetime.timedelta(days=days_to_sunday + 7)
    week_folder = os.path.join(source_root, f"Week of {week_start.strftime('%Y-%m-%d')}")
    logger.info(f"Processing weekly directory: {week_folder}")

    if not os.path.isdir(week_folder):
        logger.warning(f"Weekly directory not found: '{week_folder}'. Nothing to process."); sys.exit(0)
    
    # NEW: Cache for loaded custom prompts to avoid redundant file I/O
    loaded_prompts_cache = {}

    for ext, agent_data in member_list.items():
        try:
            agent_folder = os.path.join(week_folder, ext)
            if not os.path.isdir(agent_folder): continue
            
            # --- NEW: Dynamic Prompt Selection Logic ---
            individual_prompt_to_use = prompts['individual'] # Start with the default
            custom_prompt_file = agent_data.get('prompt_file')

            if custom_prompt_file:
                logger.info(f"Agent {ext} has custom prompt specified: '{custom_prompt_file}'")
                if custom_prompt_file in loaded_prompts_cache:
                    individual_prompt_to_use = loaded_prompts_cache[custom_prompt_file]
                    logger.info(f"Using cached custom prompt '{custom_prompt_file}'.")
                else:
                    try:
                        # Construct the full path to the custom prompt file
                        custom_prompt_path = os.path.join(prompt_dir, custom_prompt_file)
                        individual_prompt_to_use = load_text_from_file(custom_prompt_path, f"Custom Prompt ({custom_prompt_file})")
                        loaded_prompts_cache[custom_prompt_file] = individual_prompt_to_use
                    except SystemExit:
                        logger.error(f"FATAL: Could not load custom prompt '{custom_prompt_file}' for agent {ext}. Skipping this agent.")
                        continue # Move to the next agent
            # --- End of New Logic ---

            output_dirs = {"docx": os.path.join(agent_folder, "AutoQA_Generated_Reports_DOCX"),
                           "json": os.path.join(agent_folder, "AutoQA_Generated_Reports_DOCX", "Individual_AI_JSON_Analyses")}
            os.makedirs(output_dirs['json'], exist_ok=True)
            
            successful_jsons = [res for path in glob.glob(os.path.join(agent_folder, "*.wav")) if (res := process_single_audio(path, output_dirs, ai_model, individual_prompt_to_use, agent_data['full_name'], week_start.strftime('%Y-%m-%d'))) is not None]

            if not successful_jsons: logger.warning(f"No successful analyses for {agent_data['full_name']}."); continue

            combined_json = perform_combined_analysis(successful_jsons, agent_data['full_name'], output_dirs['docx'], ai_model, prompts['combined'])
            if combined_json:
                docx_path = os.path.join(output_dirs['docx'], "Combined_Analysis_Report.docx")
                save_text_to_file(os.path.join(output_dirs['docx'], "Combined_Analysis_Report.json"), json.dumps(combined_json, indent=2))
                generate_combined_docx(combined_json, docx_path, agent_data['full_name'])
                
                email_json = generate_email_body(combined_json, agent_data['full_name'], ai_model, prompts['email'])
                if email_json: send_email(smtp_settings, email_settings, agent_data['email'], email_json.get('subject'), email_json.get('body'), docx_path)
                else: logger.error(f"Failed to generate email body for {agent_data['full_name']}.")
            else: logger.error(f"Failed to get combined analysis for {agent_data['full_name']}.")
        except Exception as e_agent_loop:
            logger.error(f"An unhandled error occurred while processing agent {ext}: {e_agent_loop}", exc_info=True)
    logger.info("AutoQA script execution completed.")


if __name__ == "__main__":
    cli_date, is_debug = parse_arguments()
    if not os.path.exists(CONFIG_FILE_PATH):
        print(f"FATAL: Config file '{CONFIG_FILE_PATH}' not found.", file=sys.stderr); sys.exit(1)
    config = configparser.ConfigParser(interpolation=None); config.read(CONFIG_FILE_PATH)
    try:
        main(config, cli_date, is_debug)
    except SystemExit:
        print("Script exited.")
    except Exception as e:
        traceback.print_exc(file=sys.stderr)
        if logger.hasHandlers(): logger.critical("An unhandled exception occurred in main execution.", exc_info=True)
        sys.exit(1)
